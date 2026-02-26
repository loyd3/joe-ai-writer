from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from typing import Optional
import io
import json
import tempfile
import os
from datetime import datetime
from urllib.parse import quote

from app.database import get_db
from app.api.auth import get_current_user
from app.api.projects import check_document_access
from app.models.models import Document, Project, AIMemory

# 导出库
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

router = APIRouter(prefix="/api/export", tags=["export"])


def content_disposition_attachment(filename: str) -> str:
    """生成支持中文文件名的 Content-Disposition 头（RFC 5987）。"""
    ascii_fallback = "export" + (filename[filename.rfind("."):] if "." in filename else "")
    encoded = quote(filename, safe="")
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"


@router.get("/document/{document_id}/markdown")
async def export_markdown(
    document_id: int,
    include_memory: bool = True,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """导出文档为 Markdown 格式"""
    document = check_document_access(db, document_id, current_user["id"])
    project = db.query(Project).filter(Project.id == document.project_id).first()
    
    # 生成 Markdown 内容
    md_content = generate_markdown(db, document, project, include_memory)
    
    # 创建文件流
    filename = f"{document.title}_{datetime.now().strftime('%Y%m%d')}.md"
    return StreamingResponse(
        io.StringIO(md_content),
        media_type="text/markdown",
        headers={"Content-Disposition": content_disposition_attachment(filename)}
    )


@router.get("/document/{document_id}/txt")
async def export_txt(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """导出文档为纯文本格式"""
    document = check_document_access(db, document_id, current_user["id"])
    
    # 生成纯文本内容
    text_content = generate_plain_text(document)
    
    filename = f"{document.title}_{datetime.now().strftime('%Y%m%d')}.txt"
    return StreamingResponse(
        io.StringIO(text_content),
        media_type="text/plain",
        headers={"Content-Disposition": content_disposition_attachment(filename)}
    )


@router.get("/project/{project_id}/markdown")
async def export_project_markdown(
    project_id: int,
    include_memory: bool = True,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """导出整个项目为 Markdown 格式"""
    # 检查权限
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user["id"]
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # 生成项目 Markdown
    md_content = generate_project_markdown(db, project, include_memory)
    
    filename = f"{project.title}_{datetime.now().strftime('%Y%m%d')}.md"
    return StreamingResponse(
        io.StringIO(md_content),
        media_type="text/markdown",
        headers={"Content-Disposition": content_disposition_attachment(filename)}
    )


def generate_markdown(db: Session, document: Document, project: Project, include_memory: bool) -> str:
    """生成 Markdown 格式内容"""
    lines = []
    
    # 文档标题
    lines.append(f"# {document.title}")
    lines.append("")
    lines.append(f"> 所属项目: {project.title}")
    lines.append(f"> 导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # 添加 项目设定
    if include_memory:
        memory = db.query(AIMemory).filter(AIMemory.project_id == project.id).first()
        if memory and has_memory_content(memory):
            lines.append("## 📚 项目设定")
            lines.append("")
            
            if memory.outline:
                lines.append("### 大纲")
                for item in memory.outline:
                    lines.append(f"- {item.get('title', '')}")
                lines.append("")
            
            if memory.storyline:
                lines.append("### 故事线")
                lines.append(memory.storyline)
                lines.append("")
            
            if memory.characters:
                lines.append("### 角色设定")
                for char in memory.characters:
                    lines.append(f"**{char.get('name', '')}**: {char.get('description', '')}")
                    if char.get('personality'):
                        lines.append(f"- 性格: {char['personality']}")
                    if char.get('goals'):
                        lines.append(f"- 目标: {char['goals']}")
                    lines.append("")
            
            if memory.writing_style:
                lines.append("### 写作风格")
                lines.append(memory.writing_style)
                lines.append("")
            
            lines.append("---")
            lines.append("")
    
    # 文档内容
    lines.append("## 正文")
    lines.append("")
    
    for block in (document.content or []):
        block_type = block.get("type", "paragraph")
        content = block.get("content", "")
        
        if block_type == "heading":
            lines.append(f"## {content}")
        elif block_type == "quote":
            lines.append(f"> {content}")
        elif block_type == "list":
            lines.append(f"- {content}")
        else:
            lines.append(content)
        
        lines.append("")
    
    return "\n".join(lines)

def generate_plain_text(document: Document) -> str:
    """生成纯文本格式内容"""
    lines = []
    lines.append(document.title)
    lines.append("=" * len(document.title))
    lines.append("")
    
    for block in (document.content or []):
        content = block.get("content", "")
        if content.strip():
            lines.append(content)
            lines.append("")
    
    return "\n".join(lines)

def generate_project_markdown(db: Session, project: Project, include_memory: bool) -> str:
    """生成项目 Markdown"""
    lines = []
    
    # 项目标题
    lines.append(f"# {project.title}")
    lines.append("")
    if project.description:
        lines.append(f"> {project.description}")
    lines.append(f"> 导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    lines.append("=" * 50)
    lines.append("")
    
    # 添加 项目设定
    if include_memory:
        memory = db.query(AIMemory).filter(AIMemory.project_id == project.id).first()
        if memory and has_memory_content(memory):
            lines.append("## 📚 项目设定")
            lines.append("")
            
            if memory.outline:
                lines.append("### 大纲")
                for i, item in enumerate(memory.outline, 1):
                    lines.append(f"{i}. {item.get('title', '')}")
                lines.append("")
            
            if memory.storyline:
                lines.append("### 故事线")
                lines.append(memory.storyline)
                lines.append("")
            
            if memory.characters:
                lines.append("### 角色设定")
                for char in memory.characters:
                    lines.append(f"**{char.get('name', '')}**")
                    lines.append(f"- 描述: {char.get('description', '')}")
                    if char.get('personality'):
                        lines.append(f"- 性格: {char['personality']}")
                    if char.get('background'):
                        lines.append(f"- 背景: {char['background']}")
                    if char.get('goals'):
                        lines.append(f"- 目标: {char['goals']}")
                    lines.append("")
            
            if memory.world_building:
                lines.append("### 世界观")
                for key, value in memory.world_building.items():
                    lines.append(f"- **{key}**: {value}")
                lines.append("")
            
            if memory.writing_style:
                lines.append("### 写作风格")
                lines.append(memory.writing_style)
                lines.append("")
            
            if memory.key_points:
                lines.append("### 关键情节点")
                for point in memory.key_points:
                    lines.append(f"- {point}")
                lines.append("")
            
            lines.append("=" * 50)
            lines.append("")
    
    # 添加所有文档
    documents = db.query(Document).filter(
        Document.project_id == project.id
    ).order_by(Document.order_index).all()
    
    if documents:
        lines.append(f"## 📄 文档列表 ({len(documents)} 篇)")
        lines.append("")
        
        for i, doc in enumerate(documents, 1):
            lines.append(f"### {i}. {doc.title}")
            lines.append("")
            
            for block in (doc.content or []):
                block_type = block.get("type", "paragraph")
                content = block.get("content", "")
                
                if block_type == "heading":
                    lines.append(f"**{content}**")
                elif block_type == "quote":
                    lines.append(f"> {content}")
                elif block_type == "list":
                    lines.append(f"- {content}")
                else:
                    lines.append(content)
                
                lines.append("")
            
            lines.append("---")
            lines.append("")
    
    return "\n".join(lines)

def has_memory_content(memory: AIMemory) -> bool:
    """检查是否有项目设定内容"""
    return bool(
        memory.outline or
        memory.storyline or
        memory.characters or
        memory.world_building or
        memory.writing_style or
        memory.key_points
    )


@router.get("/document/{document_id}/pdf")
async def export_pdf(
    document_id: int,
    include_memory: bool = True,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """导出文档为 PDF 格式"""
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF 导出功能不可用，请安装 reportlab")
    
    document = check_document_access(db, document_id, current_user["id"])
    project = db.query(Project).filter(Project.id == document.project_id).first()
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name
    
    try:
        # 创建 PDF
        doc = SimpleDocTemplate(
            tmp_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # 添加标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=WD_ALIGN_PARAGRAPH.CENTER if hasattr(WD_ALIGN_PARAGRAPH, 'CENTER') else 1
        )
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        # 添加项目设定
        if include_memory:
            memory = db.query(AIMemory).filter(AIMemory.project_id == project.id).first()
            if memory and has_memory_content(memory):
                story.append(Paragraph("项目设定", styles['Heading2']))
                story.append(Spacer(1, 12))
                
                if memory.outline:
                    story.append(Paragraph("大纲", styles['Heading3']))
                    for item in memory.outline:
                        story.append(Paragraph(f"• {item.get('title', '')}", styles['BodyText']))
                    story.append(Spacer(1, 12))
                
                story.append(PageBreak())
        
        # 添加文档内容
        story.append(Paragraph("正文", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        for block in (document.content or []):
            block_type = block.get("type", "paragraph")
            content = block.get("content", "")
            
            if not content.strip():
                continue
                
            if block_type == "heading":
                story.append(Paragraph(content, styles['Heading2']))
            elif block_type == "subheading":
                story.append(Paragraph(content, styles['Heading3']))
            elif block_type == "quote":
                quote_style = ParagraphStyle(
                    'Quote',
                    parent=styles['BodyText'],
                    leftIndent=20,
                    rightIndent=20,
                    textColor='grey'
                )
                story.append(Paragraph(f'"{content}"', quote_style))
            elif block_type == "code":
                code_style = ParagraphStyle(
                    'Code',
                    parent=styles['Code'],
                    fontName='Courier',
                    fontSize=9,
                    leftIndent=20
                )
                story.append(Paragraph(content.replace('\n', '<br/>'), code_style))
            elif block_type == "divider":
                story.append(Spacer(1, 20))
            else:
                story.append(Paragraph(content, styles['BodyText']))
            
            story.append(Spacer(1, 6))
        
        doc.build(story)
        
        filename = f"{document.title}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return FileResponse(
            tmp_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": content_disposition_attachment(filename)}
        )
    except Exception as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=500, detail=f"PDF 导出失败: {str(e)}")


@router.get("/document/{document_id}/docx")
async def export_docx(
    document_id: int,
    include_memory: bool = True,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """导出文档为 Word (DOCX) 格式"""
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=503, detail="Word 导出功能不可用，请安装 python-docx")
    
    document = check_document_access(db, document_id, current_user["id"])
    project = db.query(Project).filter(Project.id == document.project_id).first()
    
    # 创建 Word 文档
    doc = DocxDocument()
    
    # 添加标题
    title = doc.add_heading(document.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if hasattr(WD_ALIGN_PARAGRAPH, 'CENTER') else 1
    
    # 添加元信息
    meta = doc.add_paragraph()
    meta.add_run(f"所属项目: {project.title}\n").italic = True
    meta.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    
    doc.add_paragraph()  # 空行
    
    # 添加项目设定
    if include_memory:
        memory = db.query(AIMemory).filter(AIMemory.project_id == project.id).first()
        if memory and has_memory_content(memory):
            doc.add_heading("项目设定", level=1)
            
            if memory.outline:
                doc.add_heading("大纲", level=2)
                for item in memory.outline:
                    doc.add_paragraph(item.get("title", ""), style='List Bullet')
            
            if memory.storyline:
                doc.add_heading("故事线", level=2)
                doc.add_paragraph(memory.storyline)
            
            if memory.characters:
                doc.add_heading("角色设定", level=2)
                for char in memory.characters:
                    p = doc.add_paragraph()
                    p.add_run(char.get("name", "")).bold = True
                    p.add_run(f": {char.get('description', '')}")
            
            doc.add_page_break()
    
    # 添加正文
    doc.add_heading("正文", level=1)
    
    for block in (document.content or []):
        block_type = block.get("type", "paragraph")
        content = block.get("content", "")
        
        if not content.strip():
            continue
        
        if block_type == "heading":
            doc.add_heading(content, level=2)
        elif block_type == "subheading":
            doc.add_heading(content, level=3)
        elif block_type == "quote":
            p = doc.add_paragraph(content)
            p.style = 'Intense Quote'
        elif block_type == "list":
            doc.add_paragraph(content, style='List Bullet')
        elif block_type == "code":
            p = doc.add_paragraph(content)
            p.style = 'No Spacing'
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(10)
        elif block_type == "divider":
            doc.add_paragraph("* * *")
        else:
            doc.add_paragraph(content)
    
    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    
    try:
        doc.save(tmp_path)
        filename = f"{document.title}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return FileResponse(
            tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            headers={"Content-Disposition": content_disposition_attachment(filename)}
        )
    except Exception as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=500, detail=f"Word 导出失败: {str(e)}")
